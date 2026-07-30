"""格式分流:把上传的图纸/文档统一成 images(JPEG bytes)+ text_map。

- PDF:PyMuPDF 渲染页面图片 + pdfplumber 提词级 bbox(归一化)。
- 图片(jpeg/png/webp):转 JPEG;text_map 空(纯视觉)。
- Word(.docx):python-docx 提文字+表格(无坐标);images 空。
- Excel(.xlsx):openpyxl 读所有 sheet 单元格;text_map page=sheet 序号。
- 其它:抛 UnsupportedFormatError。
"""

from __future__ import annotations

import io

from PIL import Image


class UnsupportedFormatError(Exception):
    """不支持的文件格式。"""


# ---- bbox 归一化 ----
def normalize_bbox(bbox, page_width, page_height) -> list[float]:
    """pdfplumber 词 bbox 为 (x0, top, x1, bottom);转成归一化 [x, y, w, h] in [0,1]。"""
    if not bbox or page_width <= 0 or page_height <= 0:
        return [0.0, 0.0, 0.0, 0.0]
    x0, top, x1, bottom = bbox
    x = max(0.0, min(1.0, x0 / page_width))
    y = max(0.0, min(1.0, top / page_height))
    w = max(0.0, min(1.0, (x1 - x0) / page_width))
    h = max(0.0, min(1.0, (bottom - top) / page_height))
    return [round(x, 5), round(y, 5), round(w, 5), round(h, 5)]


# ---- 图片 ----
def image_to_jpeg(data: bytes, content_type: str | None = None) -> bytes:
    """把任意图片字节转 JPEG。"""
    img = Image.open(io.BytesIO(data))
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    return buf.getvalue()


# ---- PDF ----
def render_pdf_pages(pdf_bytes: bytes, dpi: int = 150) -> list[bytes]:
    """用 PyMuPDF 逐页渲染成 JPEG bytes。"""
    import fitz  # PyMuPDF

    images: list[bytes] = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        for page in doc:
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            png_bytes = pix.tobytes("png")
            # 统一转 JPEG
            img = Image.open(io.BytesIO(png_bytes))
            if img.mode != "RGB":
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=90)
            images.append(buf.getvalue())
    finally:
        doc.close()
    return images


def detect_text_layer(pdf_bytes: bytes) -> bool:
    """判断 PDF 是否包含可提取文字层。"""
    import pdfplumber

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            if page.extract_words():
                return True
    return False


def extract_pdf_text(pdf_bytes: bytes) -> list[dict]:
    """用 pdfplumber 提词级 bbox,归一化到 [0,1],page 从 1 开始。"""
    import pdfplumber

    text_map: list[dict] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            width = page.width or 1
            height = page.height or 1
            for word in page.extract_words():
                text = (word.get("text") or "").strip()
                if not text:
                    continue
                bbox = normalize_bbox(
                    (word.get("x0"), word.get("top"), word.get("x1"), word.get("bottom")),
                    width,
                    height,
                )
                text_map.append({"text": text, "bbox": bbox, "page": page_idx})
    return text_map


# ---- Word ----
def extract_docx_text(docx_bytes: bytes) -> list[dict]:
    """用 python-docx 提文字 + 表格;text_map 无坐标,page=1。"""
    import docx

    doc = docx.Document(io.BytesIO(docx_bytes))
    text_map: list[dict] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            text_map.append({"text": text, "bbox": None, "page": 1})
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            joined = " | ".join(x for x in cells if x)
            if joined:
                text_map.append({"text": joined, "bbox": None, "page": 1})
    return text_map


# ---- Excel ----
def extract_xlsx_text(xlsx_bytes: bytes) -> list[dict]:
    """用 openpyxl 读所有 sheet 单元格;page=sheet 序号(从 1)。"""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    text_map: list[dict] = []
    for sheet_idx, ws in enumerate(wb.worksheets, start=1):
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell is None:
                    continue
                text = str(cell).strip()
                if text:
                    text_map.append({"text": text, "bbox": None, "page": sheet_idx})
    wb.close()
    return text_map


# ---- 分流总入口 ----
_IMAGE_CT = {"image/jpeg", "image/png", "image/webp", "image/jpg"}


def _guess_content_type(data: bytes, content_type: str | None, filename: str | None) -> str:
    if content_type:
        return content_type.lower()
    if filename:
        name = filename.lower()
        if name.endswith(".pdf"):
            return "application/pdf"
        if name.endswith((".jpg", ".jpeg")):
            return "image/jpeg"
        if name.endswith(".png"):
            return "image/png"
        if name.endswith(".webp"):
            return "image/webp"
        if name.endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if name.endswith(".xlsx"):
            return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    # magic bytes
    if data[:4] == b"%PDF":
        return "application/pdf"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:2] == b"\xff\xd8":
        return "image/jpeg"
    if data[:4] == b"PK\x03\x04":
        # docx/xlsx 都是 zip,无法仅凭 magic 区分;默认按 .docx
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return (content_type or "application/octet-stream").lower()


def format_split(
    content_type: str | None,
    data: bytes,
    filename: str | None = None,
) -> dict:
    """主分流入口。返回 {images, text_map, source_kind}。

    images: list[bytes JPEG]
    text_map: list[{text, bbox, page}]
    source_kind: 'pdf'|'image'|'docx'|'xlsx'
    """
    ct = _guess_content_type(data, content_type, filename)

    if ct == "application/pdf":
        images = render_pdf_pages(data)
        text_map = extract_pdf_text(data)
        return {"images": images, "text_map": text_map, "source_kind": "pdf"}

    if ct in _IMAGE_CT:
        images = [image_to_jpeg(data, ct)]
        return {"images": images, "text_map": [], "source_kind": "image"}

    if ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text_map = extract_docx_text(data)
        return {"images": [], "text_map": text_map, "source_kind": "docx"}

    if ct == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        text_map = extract_xlsx_text(data)
        return {"images": [], "text_map": text_map, "source_kind": "xlsx"}

    raise UnsupportedFormatError(f"不支持的文件格式:content_type={ct}, filename={filename}")
